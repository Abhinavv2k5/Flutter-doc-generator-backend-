from flask import Flask, request, send_file
from docx import Document
from docx.shared import Inches
import os
import tempfile
import uuid

app = Flask(__name__)

@app.route("/generate-report", methods=["POST"])
def generate_report():
    try:
        data = request.form.to_dict()
        template = request.files.get("template")
        if not template:
            return "Template required", 400

        doc = Document(template)

        # Mapping fields from Flutter → Word placeholders
        field_map = {
            "Replica #": "ReplicaNo",
            "Component / Location": "ComponentLocation",
            "Material of Construction": "Material",
            "Hardness In HB": "Hardness",
            "Etchant": "Etchant",
            "Microstructure": "Microstructure",
            "Structural Damage Rating": "DamageRating",
            "Life Exhaustion": "LifeExhaustion",
            "Inspection Interval": "InspectionInterval",
            "Result / Remarks": "ResultRemarks"
        }

        image_placeholders = {
            "location_photo": "{{PhotoLocation}}",
            "magnification_100x": "{{Magnification100x}}",
            "magnification_500x": "{{Magnification500x}}"
        }

        # Slightly larger image width
        IMAGE_WIDTH = Inches(2.4)

        def replace_text_paragraph(paragraph):
            full_text = ''.join(run.text for run in paragraph.runs)
            replaced = full_text
            for key, value in data.items():
                if key in field_map:
                    placeholder = f"{{{{{field_map[key]}}}}}"
                    replaced = replaced.replace(placeholder, value)
            if full_text != replaced:
                for run in paragraph.runs:
                    run.text = ''
                paragraph.runs[0].text = replaced

        # Replace text in paragraphs
        for para in doc.paragraphs:
            replace_text_paragraph(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_text_paragraph(para)

        # Replace image placeholders
        def replace_image_paragraphs(paragraphs):
            for para in paragraphs:
                for field, placeholder in image_placeholders.items():
                    if placeholder in para.text and field in request.files:
                        para.clear()
                        para.add_run().add_picture(request.files[field], width=IMAGE_WIDTH)

        replace_image_paragraphs(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_image_paragraphs(cell.paragraphs)

        # Custom filename support
        user_filename = data.get("filename", "generated_report").strip()
        if not user_filename.lower().endswith(".docx"):
            user_filename += ".docx"

        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{user_filename}")
        doc.save(temp_path)

        return send_file(temp_path, as_attachment=True, download_name=user_filename)

    except Exception as e:
        app.logger.error(f"Server error: {e}")
        return f"Server error: {e}", 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
